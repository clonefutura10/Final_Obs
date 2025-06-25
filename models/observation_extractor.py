import requests
import base64
import json
import google.generativeai as genai
import docx
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import time
import re
from datetime import datetime
from config import Config
import os

class ObservationExtractor:
    def __init__(self):
        self.ocr_api_key = Config.OCR_API_KEY
        self.groq_api_key = Config.GROQ_API_KEY
        self.gemini_api_key = Config.GOOGLE_API_KEY
        genai.configure(api_key=self.gemini_api_key)

    def image_to_base64(self, image_file):
        """Convert image file to base64 string"""
        return base64.b64encode(image_file.read()).decode('utf-8')

    def extract_text_with_ocr(self, image_file):
        """Extract text from image using OCR.space API"""
        try:
            # Get file extension
            file_type = image_file.filename.split('.')[-1].lower()
            if file_type == 'jpeg':
                file_type = 'jpg'

            # Convert image to base64
            image_file.seek(0)  # Reset file pointer
            base64_image = self.image_to_base64(image_file)
            base64_image_with_prefix = f"data:image/{file_type};base64,{base64_image}"

            # Prepare request payload
            payload = {
                'apikey': self.ocr_api_key,
                'language': 'eng',
                'isOverlayRequired': False,
                'iscreatesearchablepdf': False,
                'issearchablepdfhidetextlayer': False,
                'OCREngine': 2,
                'detectOrientation': True,
                'scale': True,
                'base64Image': base64_image_with_prefix
            }

            # Send request to OCR API
            response = requests.post(
                'https://api.ocr.space/parse/image',
                data=payload,
                headers={'apikey': self.ocr_api_key}
            )

            response.raise_for_status()
            data = response.json()

            # Process response
            if not data.get('ParsedResults') or len(data['ParsedResults']) == 0:
                error_msg = data.get('ErrorMessage', 'No parsed results returned')
                raise Exception(f"OCR Error: {error_msg}")

            parsed_result = data['ParsedResults'][0]
            if parsed_result.get('ErrorMessage'):
                raise Exception(f"OCR Error: {parsed_result['ErrorMessage']}")

            extracted_text = parsed_result['ParsedText']

            if not extracted_text or not extracted_text.strip():
                raise Exception("No text was detected in the image")

            return extracted_text

        except Exception as e:
            raise Exception(f"OCR Error: {str(e)}")

    def process_with_groq(self, extracted_text):
        """Process extracted text with Groq AI"""
        try:
            system_prompt = """You are an AI assistant for a learning observation system. Extract and structure information from the provided observation sheet text.

CRITICAL INSTRUCTIONS FOR NAME HANDLING:
- Do NOT use any names that appear in the observation text or audio transcription
- For the "studentName" field, ONLY use names that are explicitly provided by the system/database
- If no database name is provided, use "Student" as the default
- NEVER assume gender - always refer to the student as "the student" or "Student" in descriptions
- Do not use gender-specific pronouns (he/his, she/her) in any part of the response

The observation sheets typically have the following structure:
- Title (usually "The Observer")
- Student information (Name, Roll Number/ID) - IGNORE names from this section
- Date and Time information
- Core Observation Section with time slots
- Teaching content for each time slot
- Learning details (what was learned, tools used, etc.)

Format your response as JSON with the following structure:
{
  "studentName": "Use ONLY the database-provided name, never from observation text",
  "studentId": "Student ID or Roll Number",
  "className": "Class name or subject being taught",
  "date": "Date of observation",
  "observations": "Detailed description of what was learned - refer to 'the student' not by name",
  "strengths": ["List of strengths observed - use 'the student' in descriptions"],
  "areasOfDevelopment": ["List of areas where the student needs improvement - use 'the student'"],
  "recommendations": ["List of recommended actions - refer to 'the student'"],
  "themeOfDay": "Main theme or topic of the day",
  "curiositySeed": "Something that sparked the child's interest"
}

For observations, provide full detailed descriptions like:
"The student learned how to make maggi from their parent through in-person mode, including all steps from boiling water to adding spices"

IMPORTANT: Never use gender-specific language or names from the observation text. Always refer to 'the student' in descriptions."""

            # Send request to Groq API
            response = requests.post(
                'https://api.groq.com/openai/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {self.groq_api_key}',
                    'Content-Type': 'application/json'
                },
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {
                            "role": "system",
                            "content": system_prompt
                        },
                        {
                            "role": "user",
                            "content": f"Extract and structure the following text from an observation sheet: {extracted_text}"
                        }
                    ],
                    "temperature": 0.2,
                    "response_format": {"type": "json_object"}
                }
            )

            response.raise_for_status()
            data = response.json()

            # Extract the JSON content
            ai_response = data['choices'][0]['message']['content']
            return json.loads(ai_response)

        except Exception as e:
            raise Exception(f"Groq API Error: {str(e)}")

    def transcribe_with_assemblyai(self, audio_file):
        """Transcribe audio using AssemblyAI API"""
        if not Config.ASSEMBLYAI_API_KEY:
            return "Error: AssemblyAI API key is not configured."

        headers = {
            "authorization": Config.ASSEMBLYAI_API_KEY,
            "content-type": "application/json"
        }

        try:
            # Upload the audio file
            audio_file.seek(0)  # Reset file pointer
            upload_response = requests.post(
                "https://api.assemblyai.com/v2/upload",
                headers={"authorization": Config.ASSEMBLYAI_API_KEY},
                data=audio_file.read()
            )

            if upload_response.status_code != 200:
                return f"Error uploading audio: {upload_response.text}"

            upload_url = upload_response.json()["upload_url"]

            # Request transcription
            transcript_request = {
                "audio_url": upload_url,
                "language_code": "en"
            }

            transcript_response = requests.post(
                "https://api.assemblyai.com/v2/transcript",
                json=transcript_request,
                headers=headers
            )

            if transcript_response.status_code != 200:
                return f"Error requesting transcription: {transcript_response.text}"

            transcript_id = transcript_response.json()["id"]

            # Poll for completion
            status = "processing"
            while status != "completed" and status != "error":
                polling_response = requests.get(
                    f"https://api.assemblyai.com/v2/transcript/{transcript_id}",
                    headers=headers
                )

                if polling_response.status_code != 200:
                    return f"Error checking transcription status: {polling_response.text}"

                polling_data = polling_response.json()
                status = polling_data["status"]

                if status == "completed":
                    return polling_data["text"]
                elif status == "error":
                    return f"Transcription error: {polling_data.get('error', 'Unknown error')}"

                time.sleep(2)

            return "Error: Transcription timed out or failed."
        except Exception as e:
            return f"Error during transcription: {str(e)}"

    def generate_conversational_transcript(self, raw_text):
        """Convert raw observations/transcript into conversational format using Gemini API"""
        try:
            prompt = f"""
            Convert the following observation text into a conversational format between an Observer and a Child. 

CRITICAL INSTRUCTIONS:
- NEVER use any names that appear in the raw text or audio
- Always refer to the child as "Child" in the dialogue labels
- Do not assume gender - avoid using he/his, she/her pronouns
- Use gender-neutral language throughout the conversation

Format it as a natural dialogue where:
            - Observer speaks first with questions, instructions, or observations
            - Child responds naturally based on the context
            - Use "Observer:" and "Child:" labels for each speaker
            - Make it realistic and age-appropriate
            - If the text is already conversational, just format it properly
            - If it's narrative observations, convert them into likely dialogue
            - Create a natural flow of conversation that would lead to the observations described
            - Include educational moments and learning interactions
            - NEVER use names from the original text - always use "Child:" as the label
            - Avoid gender assumptions in the dialogue content

            Original observation text:
            {raw_text}

            Please format as a realistic conversation:
            Observer: [what the observer might have said or asked]
            Child: [how the child might have responded based on the context]
            Observer: [follow-up from observer]
            Child: [child's response]

            Keep it natural, educational, and age-appropriate. Make sure the conversation flows logically and would realistically result in the observations described in the original text. Remember to never use names from the original text and avoid gender-specific language.
            """

            # Use the same Gemini API pattern as your existing methods
            model = genai.GenerativeModel('gemini-2.0-flash')
            response = model.generate_content([
                {"role": "user", "parts": [{"text": prompt}]}
            ])

            if response and response.text:
                return response.text.strip()
            else:
                return self._basic_transcript_formatting(raw_text)

        except Exception as e:
            # Fallback to basic formatting if API fails
            return self._basic_transcript_formatting(raw_text)

    def _basic_transcript_formatting(self, raw_text):
        """Basic fallback transcript formatting"""
        lines = raw_text.split('\n')
        formatted_lines = []

        for i, line in enumerate(lines):
            line = line.strip()
            if line:
                if i % 2 == 0:
                    formatted_lines.append(f"Observer: {line}")
                else:
                    formatted_lines.append(f"Child: {line}")

        if not formatted_lines:
            # If no lines, create a basic conversation
            formatted_lines = [
                "Observer: Can you tell me about what you learned today?",
                f"Child: {raw_text[:100]}..." if len(raw_text) > 100 else f"Child: {raw_text}",
                "Observer: That's wonderful! Can you tell me more about it?",
                "Child: Yes, I enjoyed learning about this topic."
            ]

        return '\n'.join(formatted_lines)

    def generate_report_from_text(self, text_content, user_info):
        """Generate a structured report from text using Google Gemini"""
        prompt = f"""
        You are an educational observer tasked with generating a comprehensive and accurate Daily Growth Report based on the following observational notes from a student session. 

        CRITICAL INSTRUCTIONS FOR NAME AND GENDER USAGE:
        - NEVER extract or use any name from the audio transcription or text content
        - ALWAYS use the exact name provided: {user_info['student_name']}
        - NEVER assume gender - always refer to the student by their name "{user_info['student_name']}" throughout the report
        - Do not use pronouns like he/his, she/her, they/them - use the student's name consistently
        - If you need to refer to the student multiple times, use "{user_info['student_name']}" or "the student"

        Please carefully analyze the given text and complete the report using the exact format, emojis, section titles, and scoring rubrics as described below. The student should be referred to consistently using their provided name "{user_info['student_name']}" - never use gender-specific pronouns or names from the audio/text content.

        📌 Important Instructions for the Report:
        - Follow the format exactly as shown below.
        - Make reasonable inferences for items not explicitly stated in the text.
        - Ensure that the final Overall Growth Score and category (🟢/💚/⚠️/📈) accurately reflects the number of active areas, according to:
        🟢 Excellent (7/7 areas) – Clear growth with strong evidence
        💚 Good (5-6 areas) – Solid engagement with positive trends        
        ⚠️ Fair (3-4 areas) – Some engagement, needs encouragement        
        📈 Needs Work (1-2 areas) – Area not activated or underperforming today
        - Include the new Communication Skills & Thought Clarity section.
        - The tone should be professional, warm, and insightful — aimed at helping parents understand their child's daily growth.
        - REMEMBER: Always use "{user_info['student_name']}" instead of any pronouns or names from the content
        
        Instructions for Report Generation
        Assign scores based on clear, evidence-backed observations for each area.

        Explain each score with a specific reason—avoid generalizations or repeated points. Every score must be justified individually and precisely.

        Use the following rating scale consistently:

        Ratings Scale:
         Excellent (7/7 areas) – Clear growth with strong evidence
         Good (5-6 areas) – Solid engagement with positive trends        
         Fair (3-4 areas) – Some engagement, needs encouragement        
         Needs Work (1-2 areas) – Area not activated or underperforming today
        Always include the complete legend in every report so the evaluator or reader can cross-check scores against the criteria.

        Ensure the entire report strictly follows the legend and that scoring aligns accurately with the defined scale.

        Do not use tables for the "Growth Metrics & Observations" section. Present the content in a well-spaced, structured paragraph format to preserve formatting integrity across platforms.
        📝 TEXT CONTENT:
        {text_content}

        🧾 Daily Growth Report Format for Parents

        🧒 Child's Name: {user_info['student_name']}
        📅 Date: [{user_info['session_date']}]
        🌱 Curiosity Seed Explored: [Extract from text]

        📊 Growth Metrics & Observations
        Growth Area | Rating | Observation Summary
        🧠 Intellectual | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        😊 Emotional | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🤝 Social | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🎨 Creativity | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🏃 Physical | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🧭 Character/Values | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]
        🚀 Planning/Independence | [✅ Excellent/✅ Good/⚠️ Fair/📈 Needs Work] | [Brief summary]

        🌈 Curiosity Response Index: [1-10] / 10  
        [Brief explanation of {user_info['student_name']}'s engagement with the curiosity seed]

        🗣️ Communication Skills & Thought Clarity
        • Confidence level: [Describe based on speech and tone in text]  
        • Clarity of thought: [Describe {user_info['student_name']}'s ability to express thoughts clearly and independently]  
        • Participation & engagement: [Describe based on frequency and quality of responses]  
        • Sequence of explanation: [Describe structure and coherence of thought process]  

        🧠 Overall Growth Score:  
        [🔵 Balanced Growth / 🟡 Moderate Growth / 🔴 Limited Growth] – [X/7] Areas Active 
        [Brief recommendation for next steps or continued development for {user_info['student_name']}]

        📣 Note for Parent:  
        [Comprehensive summary for parents with actionable insights and encouragement based on today's session for {user_info['student_name']}]

        🟢 Legend

        ✅ Performance by Area
        🟢 Excellent (7/7 areas) – Clear growth with strong evidence
        💚 Good (5-6 areas) – Solid engagement with positive trends        
        ⚠️ Fair (3-4 areas) – Some engagement, needs encouragement        
        📈 Needs Work (1-2 areas) – Area not activated or underperforming today

        give the entire report such that its a direct send worthy item, so all things should always be there and no other unecessary words in the response. No repetation.
        """

        try:
            model = genai.GenerativeModel('gemini-2.0-flash')
            response = model.generate_content([
                {"role": "user", "parts": [{"text": prompt}]}
            ])
            return response.text
        except Exception as e:
            return f"Error generating report: {str(e)}"

    def create_word_document_with_emojis(self, report_content):
        """Create a Word document from the report content with emoji support"""
        doc = docx.Document()

        # Set document encoding and font for emoji support
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Segoe UI Emoji'

        title = doc.add_heading('📋 Daily Growth Report', 0)
        title.runs[0].font.name = 'Segoe UI Emoji'

        # Process the report content line by line
        lines = report_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Clean up markdown formatting but preserve emojis
            line = line.replace('**', '')

            if line.startswith(('🧒', '📅', '🌱')):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.name = 'Segoe UI Emoji'
                run.font.size = docx.shared.Pt(12)
            elif line.startswith('📊'):
                heading = doc.add_heading(line, level=1)
                heading.runs[0].font.name = 'Segoe UI Emoji'
            elif line.startswith(('🧠', '😊', '🤝', '🎨', '🏃', '🧭', '🚀')):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.name = 'Segoe UI Emoji'
                run.font.size = docx.shared.Pt(11)
            elif line.startswith(('🌈', '🗣️')):
                heading = doc.add_heading(line, level=2)
                heading.runs[0].font.name = 'Segoe UI Emoji'
            elif line.startswith('🧠 Overall'):
                heading = doc.add_heading(line, level=2)
                heading.runs[0].font.name = 'Segoe UI Emoji'
            elif line.startswith('📣'):
                heading = doc.add_heading(line, level=2)
                heading.runs[0].font.name = 'Segoe UI Emoji'
            elif line.startswith('🟢 Legend'):
                heading = doc.add_heading(line, level=3)
                heading.runs[0].font.name = 'Segoe UI Emoji'
            elif line.startswith(('✅', '⚠️', '📈', '🔵', '🟢', '🟡', '🔴', '•', '💚')):
                p = doc.add_paragraph(line, style='List Bullet')
                p.runs[0].font.name = 'Segoe UI Emoji'
                p.runs[0].font.size = docx.shared.Pt(10)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Segoe UI Emoji'
                run.font.size = docx.shared.Pt(10)

        # Save to BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        return docx_bytes

    def create_pdf_alternative(self, report_content):
        """Create PDF using reportlab instead of WeasyPrint"""
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import black, blue, green, red

        # Create emoji mapping for text replacement
        emoji_map = {
            '🧒': '[Child]',
            '📅': '[Date]',
            '🌱': '[Curiosity Seed]',
            '📊': '[Growth Metrics]',
            '🧠': '[Intellectual]',
            '😊': '[Emotional]',
            '🤝': '[Social]',
            '🎨': '[Creativity]',
            '🏃': '[Physical]',
            '🧭': '[Character/Values]',
            '🚀': '[Planning/Independence]',
            '🌈': '[Curiosity Response]',
            '🗣️': '[Communication Skills]',
            '📣': '[Note for Parent]',
            '🟢': '[Excellent]',
            '✅': '[Good]',
            '⚠️': '[Fair]',
            '📈': '[Needs Work]',
            '🔵': '[Balanced Growth]',
            '🟡': '[Moderate Growth]',
            '🔴': '[Limited Growth]',
            '💚': '[Good Score]',
            '📋': '[Report]'
        }

        # Replace emojis with readable text
        pdf_content = report_content
        for emoji, replacement in emoji_map.items():
            pdf_content = pdf_content.replace(emoji, replacement)

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        styles = getSampleStyleSheet()

        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
            textColor=blue,
            alignment=1  # Center alignment
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=10,
            textColor=black,
            fontName='Helvetica-Bold'
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            fontName='Helvetica'
        )

        story = []

        # Add title
        title = Paragraph("[Report] Daily Growth Report", title_style)
        story.append(title)
        story.append(Spacer(1, 12))

        # Process report content
        lines = pdf_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            try:
                if line.startswith(('[Child]', '[Date]', '[Curiosity Seed]')):
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                elif line.startswith('[Growth Metrics]'):
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                elif line.startswith(('[Intellectual]', '[Emotional]', '[Social]', '[Creativity]', '[Physical]',
                                      '[Character/Values]', '[Planning/Independence]')):
                    para = Paragraph(f"<b>{line}</b>", normal_style)
                    story.append(para)
                elif line.startswith(('[Curiosity Response]', '[Communication Skills]', '[Note for Parent]')):
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                elif 'Overall Growth Score' in line:
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                elif line.startswith('[Excellent] Legend'):
                    para = Paragraph(f"<b>{line}</b>", heading_style)
                    story.append(para)
                else:
                    para = Paragraph(line, normal_style)
                    story.append(para)

                story.append(Spacer(1, 4))

            except Exception as e:
                # Skip problematic lines
                continue

        doc.build(story)
        buffer.seek(0)

        return buffer

    def create_pdf_with_emojis(self, report_content):
        """Alias for create_pdf_alternative for backward compatibility"""
        return self.create_pdf_alternative(report_content)

    def create_word_document(self, report_content):
        """Legacy method - calls the emoji version"""
        return self.create_word_document_with_emojis(report_content)

    def send_email(self, recipient_email, subject, message):
        """Send email with the observation report"""
        sender_email = Config.EMAIL_USER
        sender_password = Config.EMAIL_PASSWORD

        if not sender_password:
            return False, "Email password not configured"

        smtp_server = "smtp.gmail.com"
        smtp_port = 587

        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = recipient_email
        msg["Subject"] = subject

        # Convert message to HTML format for better emoji display
        html_message = f"""
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; }}
                pre {{ white-space: pre-wrap; font-family: inherit; }}
            </style>
        </head>
        <body>
            <pre>{message}</pre>
        </body>
        </html>
        """

        msg.attach(MIMEText(html_message, "html", "utf-8"))

        try:
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
            server.quit()
            return True, f"Email sent to {recipient_email}"
        except smtplib.SMTPAuthenticationError:
            return False, "Error: Authentication failed. Check your email and password."
        except smtplib.SMTPException as e:
            return False, f"Error: Failed to send email. {e}"
        except Exception as e:
            return False, f"Error: {str(e)}"

    def generate_custom_report_from_prompt(self, prompt, child_id):
        """Generate custom report based on prompt and stored data with new JSON format"""
        from models.database import get_supabase_client

        try:
            supabase = get_supabase_client()
            # Get all processed observations for the child
            processed_data = supabase.table('processed_observations').select("*").eq('child_id', child_id).execute()
            observations = supabase.table('observations').select("*").eq('student_id', child_id).execute()

            # Get child information
            child_data = supabase.table('children').select("name").eq('id', child_id).execute().data
            child_name = child_data[0]['name'] if child_data else 'Student'

            # Combine all data
            all_data = {
                'processed_observations': processed_data.data,
                'observations': observations.data
            }

            custom_prompt = f"""
            You are an AI assistant for a learning observation system. Extract and structure information from the provided observation data based on the user's specific request.

CRITICAL INSTRUCTIONS FOR NAME AND GENDER USAGE:
- ALWAYS use the exact name from the database: {child_name}
- NEVER use any names that appear in the observation data or audio transcriptions
- NEVER assume gender - always refer to the student by their name "{child_name}" or as "the student"
- Do not use gender-specific pronouns (he/his, she/her, they/them) anywhere in the report
- When describing activities, use "{child_name}" or "the student" consistently

            Based on the following prompt and all available data for this child, generate a comprehensive custom report in the specified JSON format:

            USER PROMPT: {prompt}

            AVAILABLE DATA: {json.dumps(all_data, indent=2)}

            Format your response as JSON with the following structure:
            {{
              "studentName": "{child_name}",
              "studentId": "Custom Report ID",
              "className": "Custom Analysis Report",
              "date": "{datetime.now().strftime('%Y-%m-%d')}",
              "observations": "Detailed description combining all relevant observations that match the user's prompt - refer to '{child_name}' or 'the student'",
              "strengths": ["List of strengths observed in {child_name} based on available data"],
              "areasOfDevelopment": ["List of areas where {child_name} needs improvement"],
              "recommendations": ["List of recommended actions for {child_name} based on the prompt and data"]
            }}

            For observations, provide full detailed descriptions like:
            "{child_name} learned how to make maggi from their parent through in-person mode, including all steps from boiling water to adding spices"

            Be creative in extracting information based on context and ensure the response directly addresses the user's prompt: "{prompt}"

            REMEMBER: Always use "{child_name}" from the database, never names from observation data, and avoid gender assumptions.
            """

            model = genai.GenerativeModel('gemini-2.0-flash')
            response = model.generate_content([
                {"role": "user", "parts": [{"text": custom_prompt}]}
            ])

            # Try to parse as JSON and format nicely
            try:
                json_response = json.loads(response.text)

                # Format the JSON response into a readable report
                formatted_report = f"""
📋 Custom Report: {json_response.get('className', 'Custom Analysis Report')}

🧒 Student Name: {json_response.get('studentName', child_name)}
📅 Date: {json_response.get('date', datetime.now().strftime('%Y-%m-%d'))}
📝 Report Type: Custom Analysis

📊 Observations Summary:
{json_response.get('observations', 'No observations available')}

⭐ Strengths Identified:
{chr(10).join([f"• {strength}" for strength in json_response.get('strengths', [])])}

📈 Areas for Development:
{chr(10).join([f"• {area}" for area in json_response.get('areasOfDevelopment', [])])}

💡 Recommendations:
{chr(10).join([f"• {rec}" for rec in json_response.get('recommendations', [])])}

📋 Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
                """

                return formatted_report.strip()

            except json.JSONDecodeError:
                # If not valid JSON, return the raw response
                return response.text

        except Exception as e:
            return f"Error generating custom report: {str(e)}"

    def generate_monthly_report_json_format(self, observations, goal_progress, child_name, year, month):
        """Generate monthly summary in the new JSON format with graph recommendations"""
        try:
            import calendar

            # Prepare data for analysis
            observation_texts = []
            all_strengths = []
            all_developments = []
            all_recommendations = []

            for obs in observations:
                observation_texts.append(obs.get('observations', ''))

                # Parse strengths, developments, and recommendations
                if obs.get('strengths'):
                    try:
                        strengths = json.loads(obs['strengths']) if isinstance(obs['strengths'], str) else obs[
                            'strengths']
                        all_strengths.extend(strengths)
                    except:
                        pass

                if obs.get('areas_of_development'):
                    try:
                        developments = json.loads(obs['areas_of_development']) if isinstance(
                            obs['areas_of_development'], str) else obs['areas_of_development']
                        all_developments.extend(developments)
                    except:
                        pass

                if obs.get('recommendations'):
                    try:
                        recommendations = json.loads(obs['recommendations']) if isinstance(obs['recommendations'],
                                                                                           str) else obs[
                            'recommendations']
                        all_recommendations.extend(recommendations)
                    except:
                        pass

            # Calculate metrics for graphs
            total_observations = len(observations)
            active_goals = len([g for g in goal_progress if g.get('status') == 'active'])
            completed_goals = len([g for g in goal_progress if g.get('status') == 'achieved'])

            # Count frequency of strengths and development areas
            strength_counts = {}
            for strength in all_strengths:
                strength_counts[strength] = strength_counts.get(strength, 0) + 1

            development_counts = {}
            for dev in all_developments:
                development_counts[dev] = development_counts.get(dev, 0) + 1

            # Prepare graph data suggestions
            graph_suggestions = []

            if total_observations > 0:
                graph_suggestions.append({
                    "type": "bar_chart",
                    "title": "Observation Frequency by Week",
                    "description": f"Shows {total_observations} observations recorded throughout the month"
                })

            if strength_counts:
                graph_suggestions.append({
                    "type": "pie_chart",
                    "title": "Top Strengths Distribution",
                    "data": strength_counts,
                    "description": f"Distribution of {len(strength_counts)} different strength areas"
                })

            if development_counts:
                graph_suggestions.append({
                    "type": "horizontal_bar",
                    "title": "Development Areas Focus",
                    "data": development_counts,
                    "description": f"Areas requiring attention with frequency counts"
                })

            if goal_progress:
                graph_suggestions.append({
                    "type": "donut_chart",
                    "title": "Goal Progress Status",
                    "data": {"Active": active_goals, "Completed": completed_goals},
                    "description": f"Goal completion status: {completed_goals} completed, {active_goals} active"
                })

            # Create comprehensive prompt for JSON generation
            monthly_prompt = f"""
            You are an AI assistant for a learning observation system. Generate a comprehensive monthly report based on the provided observation data.

CRITICAL INSTRUCTIONS FOR NAME AND GENDER USAGE:
- ALWAYS use the exact name from the database: {child_name}
- NEVER use any names that appear in the observation data or audio transcriptions
- NEVER assume gender - always refer to the student by their name "{child_name}" or as "the student"
- Do not use gender-specific pronouns (he/his, she/her, they/them) anywhere in the report
- When describing activities and progress, use "{child_name}" consistently

            MONTH: {calendar.month_name[month]} {year}
            STUDENT: {child_name}
            TOTAL OBSERVATIONS: {total_observations}
            GOALS STATUS: {active_goals} active, {completed_goals} completed

            OBSERVATION DATA: {json.dumps(observation_texts[:5], indent=2)}  # Limit for prompt size
            STRENGTHS IDENTIFIED: {list(strength_counts.keys())[:10]}
            DEVELOPMENT AREAS: {list(development_counts.keys())[:10]}

            QUANTIFIABLE METRICS FOR GRAPHS:
            {json.dumps(graph_suggestions, indent=2)}

            Format your response as JSON with the following structure:
            {{
              "studentName": "{child_name}",
              "studentId": "Monthly Report ID",
              "className": "Monthly Progress Summary",
              "date": "{calendar.month_name[month]} {year}",
              "observations": "Comprehensive monthly summary combining all {total_observations} observations, highlighting key learning moments for {child_name}, progress patterns, and notable developments throughout the month",
              "strengths": ["List of top strengths observed consistently in {child_name} throughout the month"],
              "areasOfDevelopment": ["List of areas where {child_name} needs continued focus and improvement"],
              "recommendations": ["List of specific recommended actions for {child_name} for the next month based on observed patterns"],
              "monthlyMetrics": {{
                "totalObservations": {total_observations},
                "activeGoals": {active_goals},
                "completedGoals": {completed_goals},
                "topStrengths": {dict(list(strength_counts.items())[:5])},
                "developmentFocus": {dict(list(development_counts.items())[:5])}
              }},
              "suggestedGraphs": {graph_suggestions}
            }}

            For observations, provide a comprehensive monthly summary like:
            "Throughout {calendar.month_name[month]}, {child_name} demonstrated consistent growth in multiple areas. Key learning highlights include [specific examples from observations]. {child_name} showed particular strength in [areas] while developing skills in [areas]. Notable progress was observed in [specific skills/subjects]."

            REMEMBER: Always use "{child_name}" from the database, never names from observation data, and avoid all gender assumptions.
            """

            # Generate the report using AI
            model = genai.GenerativeModel('gemini-2.0-flash')
            response = model.generate_content([
                {"role": "user", "parts": [{"text": monthly_prompt}]}
            ])

            # Try to parse as JSON and format nicely
            try:
                json_response = json.loads(response.text)

                # Format the JSON response into a readable report
                formatted_report = f"""
📋 Monthly Report: {json_response.get('className', 'Monthly Progress Summary')}

🧒 Student Name: {json_response.get('studentName', child_name)}
📅 Period: {json_response.get('date', f'{calendar.month_name[month]} {year}')}

📊 Monthly Metrics:
• Total Observations: {json_response.get('monthlyMetrics', {}).get('totalObservations', total_observations)}
• Active Goals: {json_response.get('monthlyMetrics', {}).get('activeGoals', active_goals)}
• Completed Goals: {json_response.get('monthlyMetrics', {}).get('completedGoals', completed_goals)}

📝 Monthly Observations Summary:
{json_response.get('observations', 'No observations summary available')}

⭐ Strengths Observed:
{chr(10).join([f"• {strength}" for strength in json_response.get('strengths', [])])}

📈 Areas for Development:
{chr(10).join([f"• {area}" for area in json_response.get('areasOfDevelopment', [])])}

💡 Recommendations for Next Month:
{chr(10).join([f"• {rec}" for rec in json_response.get('recommendations', [])])}

📊 Suggested Visual Analytics:
{chr(10).join([f"• {graph['title']}: {graph['description']}" for graph in json_response.get('suggestedGraphs', [])])}

📋 Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
                """

                return formatted_report.strip()

            except json.JSONDecodeError:
                # If not valid JSON, return the raw response
                return response.text

        except Exception as e:
            return f"Error generating monthly summary: {str(e)}"
